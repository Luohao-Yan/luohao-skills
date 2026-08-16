# -*- coding: utf-8 -*-
"""docx_tools.py — 可复用工具:程序化读写/改 Word .doc/.docx 文档。

设计目标:封装"读 .doc/.docx → 在指定锚点插入章节/表格/段落 → 加高亮 →
转回 .doc → 重新提取校验"这条链路上踩过坑、验证过的做法。

依赖:
    pip install python-docx
    LibreOffice (soffice)  —— 仅 .doc<->.docx 互转时需要;纯 .docx 编辑无需。

非显然要点(详见同目录 references/python-docx-gotchas.md):
    - .doc 是二进制,python-docx 无法直接读,必须先转 .docx。
    - 这些政务文档没有 "Table Grid" 样式,table.style 赋值会 KeyError → 用 set_table_borders。
    - add_table 在某些 python-docx 版本必须给 width 位置参数。
    - 混合段落+表格的批量插入,逆序 addprevious 会乱序/丢内容 → 用游标法 insert_blocks_before。
    - 校验必须从最终交付的 .doc 重新转出提取,不能只信 build 脚本的 "saved"。
"""
import os
import shutil
import subprocess
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.table import Table

# ---------- LibreOffice 路径(Windows 常见位置) ----------
SOFFICE_CANDIDATES = [
    r'C:\Program Files\LibreOffice\program\soffice.exe',
    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
]


def find_soffice():
    for p in SOFFICE_CANDIDATES:
        if os.path.exists(p):
            return p
    found = shutil.which('soffice') or shutil.which('soffice.exe')
    if found:
        return found
    raise RuntimeError('未找到 LibreOffice (soffice.exe),无法转换 .doc。请安装 LibreOffice 或把它加入 PATH。')


def convert(in_path, to_fmt, out_dir, profile='lo_docxtools'):
    """LibreOffice headless 转换。to_fmt: 'docx' 或 'doc'。返回输出文件路径。

    profile: 每次用独立的 UserInstallation profile,避免并发/缓存冲突。
    """
    soffice = find_soffice()
    os.makedirs(out_dir, exist_ok=True)
    profile_uri = 'file:///C:/Users/' + os.environ.get('USERNAME', 'Public') + \
                  '/AppData/Local/Temp/' + profile
    cmd = [soffice, f'-env:UserInstallation={profile_uri}', '--headless',
           '--convert-to', to_fmt, '--outdir', out_dir, in_path]
    r = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if r.returncode != 0:
        raise RuntimeError(f'转换失败({to_fmt}): {r.stderr or r.stdout}')
    base = os.path.splitext(os.path.basename(in_path))[0]
    out = os.path.join(out_dir, base + '.' + to_fmt)
    if not os.path.exists(out):
        raise RuntimeError(f'转换后未找到输出: {out}\nstdout={r.stdout}\nstderr={r.stderr}')
    return out


# ---------- 结构提取 ----------
def heading_level(p):
    """返回标题级别 1-9,非标题返回 0。兼容 'Heading N' 与 '标题 N'。"""
    name = (p.style.name or '') if p.style else ''
    for i in range(1, 10):
        if name in (f'Heading {i}', f'标题 {i}'):
            return i
    return 0


def iter_blocks(doc):
    """按文档真实顺序遍历段落与表格,产出 ('p', Paragraph) / ('t', Table)。"""
    body = doc.element.body
    for ch in body.iterchildren():
        if ch.tag == qn('w:p'):
            yield ('p', Paragraph(ch, doc))
        elif ch.tag == qn('w:tbl'):
            yield ('t', Table(ch, doc))


def extract_outline(doc_or_path):
    """提取标题大纲。doc_or_path 可是 Document 对象或文件路径。
    返回 [(level, text), ...]。"""
    doc = doc_or_path if isinstance(doc_or_path, Document) else Document(doc_or_path)
    out = []
    for p in doc.paragraphs:
        lvl = heading_level(p)
        t = (p.text or '').strip()
        if lvl and t:
            out.append((lvl, t))
    return out


# ---------- 锚点查找 ----------
def find_para_exact(doc, exact):
    """找文本完全等于 exact 的段落(去首尾空白)。返回 Paragraph 或 None。"""
    for p in doc.paragraphs:
        if (p.text or '').strip() == exact:
            return p
    return None


def find_para_contains(doc, fragment):
    """找文本包含 fragment 的第一个段落。"""
    for p in doc.paragraphs:
        if fragment in (p.text or ''):
            return p
    return None


# ---------- 表格边框(绕开缺失的 Table Grid 样式) ----------
def set_table_borders(table, color='000000', sz='4'):
    """给表格加全边框(单线)。不依赖任何命名样式——很多政务 .doc 没定义 Table Grid。"""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


# ---------- 高亮 ----------
def highlight_runs(paragraph, color=WD_COLOR_INDEX.YELLOW):
    for r in paragraph.runs:
        r.font.highlight_color = color


def highlight_table(table, color=WD_COLOR_INDEX.YELLOW):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.highlight_color = color


# ---------- block 定义与插入 ----------
# block 元组约定:
#   ('h', level, text)        标题,用 Heading {level} 样式
#   ('p', text) / ('note', text)   正文(Normal)
#   ('table', rows)            rows = list[list[str]], 首行当表头(加粗)

def _make_block(doc, b):
    """创建一个 block 元素(临时挂在文档末尾),返回其 XML 元素。"""
    kind = b[0]
    if kind == 'h':
        level, text = b[1], b[2]
        p = doc.add_paragraph(style=f'Heading {level}')
        p.add_run(text)
        return p._element
    if kind in ('p', 'note'):
        p = doc.add_paragraph(style='Normal')
        if b[1]:
            p.add_run(b[1])
        return p._element
    if kind == 'table':
        rows = b[1]
        # width 必传:某些 python-docx 版本无默认值,缺失会 TypeError
        t = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.5))
        set_table_borders(t)
        for i, row in enumerate(rows):
            for j, txt in enumerate(row):
                cell = t.cell(i, j)
                cell.text = str(txt)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        return t._element
    raise ValueError(f'未知 block 类型: {kind}')


def insert_blocks_before(anchor_para, blocks, highlight=False):
    """在 anchor_para 之前按正序插入 blocks。

    ⚠️ 关键:不要用"逆序 addprevious"——混合段落+表格时会乱序、丢内容。
    用游标法:首个 addprevious(anchor),之后每个 prev.addnext(el),prev=el。
    """
    doc = anchor_para._parent
    created = [_make_block(doc, b) for b in blocks]
    if highlight:
        _highlight_elements(doc, created, blocks)
    if not created:
        return created
    anchor = anchor_para._element
    anchor.addprevious(created[0])
    prev = created[0]
    for el in created[1:]:
        prev.addnext(el)
        prev = el
    return created


def append_blocks(doc, blocks, highlight=False):
    """把 blocks 按正序追加到文档末尾。add_paragraph/add_table 已自动追加到 sectPr 前。"""
    created = []
    for b in blocks:
        el = _make_block(doc, b)
        created.append(el)
    if highlight:
        _highlight_elements(doc, created, blocks)
    return created


def _highlight_elements(doc, elements, blocks):
    """对已插入的 blocks 加高亮(按 blocks 顺序对应 elements)。"""
    for el, b in zip(elements, blocks):
        kind = b[0]
        if kind == 'h' or kind in ('p', 'note'):
            highlight_runs(Paragraph(el, doc))
        elif kind == 'table':
            highlight_table(Table(el, doc))


# ---------- 校验 ----------
def has_text(doc, frag, in_tables=True):
    """doc 中是否含 frag(段落 + 可选表格)。"""
    for p in doc.paragraphs:
        if frag in (p.text or ''):
            return True
    if in_tables:
        for tb in doc.tables:
            for r in tb.rows:
                for c in r.cells:
                    if frag in (c.text or ''):
                        return True
    return False


def verify_landed(final_doc_path, expected_fragments, work_dir):
    """端到端校验:把最终交付的 .doc 重新转成 .docx,确认每个片段真实落地。
    这是本工具链的"非协商"环节——build 脚本说 saved 不代表内容在交付物里。
    返回 (ok: bool, missing: list[str])。"""
    ext = os.path.splitext(final_doc_path)[1].lower()
    if ext == '.doc':
        docx_path = convert(final_doc_path, 'docx', work_dir, profile='lo_verify')
    else:
        docx_path = final_doc_path
    d = Document(docx_path)
    missing = [f for f in expected_fragments if not has_text(d, f)]
    return (len(missing) == 0, missing)
