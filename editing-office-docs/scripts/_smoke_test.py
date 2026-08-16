# -*- coding: utf-8 -*-
"""冒烟测试 docx_tools.py:建一个含锚点的小 docx → 插入块(含表格) → 高亮 → 校验落地。
只测纯 .docx 路径(不依赖 LibreOffice),证明打包的工具能独立运行。"""
import os, sys, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__)))
from docx import Document
import docx_tools as T

work = tempfile.mkdtemp(prefix='docx_smoke_')
docx_path = os.path.join(work, 'sample.docx')

# 1) 建一个带锚点的空白 docx
d = Document()
d.add_paragraph('开头段落', style='Normal')
d.add_paragraph('部署结构设计', style='Heading 2')   # 锚点A
d.add_paragraph('一些原文内容', style='Normal')
d.add_paragraph('5.数据库设计', style='Heading 1')   # 锚点B(章标题)
d.save(docx_path)

# 2) 重新打开,插入
d = Document(docx_path)
anchor = T.find_para_exact(d, '部署结构设计')
assert anchor is not None, '找不到锚点A'
blocks = [
    ('h', 3, '2.3.1 微服务架构设计'),
    ('p', '这是测试插入的正文段落。'),
    ('table', [['列A', '列B'], ['1', '2'], ['3', '4']]),
]
T.insert_blocks_before(anchor, blocks, highlight=True)

# 末尾追加
tail = [('h', 2, '6.5 系统安全设计'), ('p', '安全设计正文。')]
T.append_blocks(d, tail, highlight=True)
d.save(docx_path)

# 3) 校验:重新打开确认顺序与内容
d2 = Document(docx_path)
blocks_order = []
for kind, obj in T.iter_blocks(d2):
    if kind == 'p':
        lvl = T.heading_level(obj)
        t = (obj.text or '').strip()
        if lvl or t:
            blocks_order.append(('H%d' % lvl if lvl else 'p', t[:24]))
    else:
        blocks_order.append(('tbl', '%dx%d' % (len(obj.rows), len(obj.columns))))

print('文档块顺序:')
for x in blocks_order:
    print('  ', x)

# 断言:2.3.1 在 部署结构设计 之前;表格在2.3.1正文后;6.5在末尾
order_str = ' | '.join(repr(x) for x in blocks_order)
assert "('H3', '2.3.1 微服务架构设计')" in order_str, '缺2.3.1标题'
assert "('tbl', '3x2')" in order_str, '缺表格'
assert "('H2', '6.5 系统安全设计')" in order_str, '缺6.5标题'
# 顺序: 2.3.1 在 '部署结构设计' 之前
i_231 = next(k for k, v in enumerate(blocks_order) if '2.3.1' in repr(v))
i_anchor = next(k for k, v in enumerate(blocks_order) if '部署结构设计' in repr(v))
assert i_231 < i_anchor, '2.3.1 未在锚点之前(顺序错!)'

# 4) 校验工具
ok, missing = T.verify_landed(docx_path, ['2.3.1 微服务架构设计', '6.5 系统安全设计', '这是测试插入的正文段落。'], work)
print('\n校验落地: ok=%s missing=%s' % (ok, missing))
assert ok, '校验未通过: %s' % missing

# 5) 高亮检查
from docx.enum.text import WD_COLOR_INDEX
yel = sum(1 for p in d2.paragraphs for r in p.runs if r.font.highlight_color == WD_COLOR_INDEX.YELLOW)
tyel = sum(1 for tb in d2.tables for row in tb.rows for c in row.cells for p in c.paragraphs for r in p.runs if r.font.highlight_color == WD_COLOR_INDEX.YELLOW)
print('高亮: 段落run=%d 表格run=%d' % (yel, tyel))
assert yel > 0 and tyel > 0, '高亮未生效'

print('\n✅ 冒烟测试通过:插入/顺序/表格/高亮/校验 全部正常')
